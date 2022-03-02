import json
import logging
import os
import re
from dataclasses import dataclass
from datetime import datetime
from functools import lru_cache
from typing import List, Optional

import docx
from docx.table import Table
from tqdm import tqdm

from oneinamillion.common import filter_word_docs
from oneinamillion.resources import PCC_PT_RECORD_RAW_DIR, PCC_PT_RECORD_DIR


@dataclass
class PatientRecord:
    date: datetime = None
    gp: str = None
    additional: str = None
    allergy: str = None
    assessment: str = None
    comment: str = None
    document: str = None
    examination: str = None
    family_history: str = None
    follow_up: str = None
    history: str = None
    lab_results: str = None
    medication: str = None
    problem: str = None
    procedure: str = None
    referral: str = None
    regime_review: str = None
    result: str = None
    social: str = None
    template_entry: str = None
    test_request: str = None
    x_ray: str = None


@dataclass
class PatientRecords:
    id: str
    records: List[PatientRecord] = None


# noinspection PyTypeChecker
class PatientRecordParser:
    doc_ids = []

    def __init__(self, from_raw=False):
        if from_raw:
            self._prepare_raw()
        self._get_doc_ids()

    def clear_cache(self):
        self._prepare_raw()
        self._get_doc_ids()

    # noinspection PyTypeChecker
    @staticmethod
    def _parse_table(table: docx.table.Table):
        if len(table.columns) < 3 or len(table.columns) > 4:
            raise NotImplementedError("Does not support current table")

        def get_date(text: str) -> Optional[datetime]:
            date = None
            try:
                date = datetime.strptime(text, "%d-%b-%Y").date()
            finally:
                return date

        def get_second_col():
            for j in range(len(table.rows)):
                first_col = table.cell(j, 0).text
                if first_col:
                    for i in range(1, len(table.columns)):
                        other_col = table.cell(j, i).text
                        if other_col and other_col != first_col:
                            return i

        records = []
        current_record = None
        current_heading = None
        idx = get_second_col()
        for ii, row in enumerate(table.rows):
            first = row.cells[0]
            third = row.cells[idx]
            first_text: str = first.text.strip()
            third_text: str = third.text.strip()
            # print(f"{first_text}\t{third_text}")
            if first_text == 'Date' and ii == 0:
                continue
            elif first_text or third_text:
                if get_date(first_text):  # if current row is a new entry
                    if current_record:
                        records.append(current_record)
                    current_record = PatientRecord()
                    current_record.date = get_date(first_text)
                    current_record.gp = third_text
                    current_heading = 'gp'
                else:
                    if first_text:
                        current_heading = first_text.replace('-', '_').replace(' ', '_').lower()
                    if third_text and current_heading:  # set attribute for current record
                        if hasattr(current_record, current_heading):
                            attr_val = getattr(current_record, current_heading)
                            if attr_val:
                                setattr(current_record, current_heading, f"{attr_val} {third_text}")
                            else:
                                setattr(current_record, current_heading, third_text)
                        else:
                            print(current_heading)
                            # raise ValueError(f"record does not support attribute {current_heading}.")
        if current_record:
            records.append(current_record)

        return records

    def _prepare_raw(self):
        logging.info("Parsing text from patient record documents")
        if not os.path.exists(PCC_PT_RECORD_RAW_DIR):
            raise FileNotFoundError
        word_docs = filter_word_docs(os.listdir(PCC_PT_RECORD_RAW_DIR))

        for doc_name in tqdm(word_docs):

            _id = re.search("(.+?)_pt_record_view", doc_name)
            if _id:
                _id = _id.group(1)

            word_doc = docx.Document(os.path.join(PCC_PT_RECORD_RAW_DIR, doc_name))
            if len(word_doc.tables) != 1:
                logging.warning(f"{doc_name} has multiple tables, must contain only one table for GP record")
                continue
                # raise ValueError(f"{doc_name} has multiple tables, must contain only one table for GP record")
            _records = self._parse_table(word_doc.tables[0])

            pt_records = PatientRecords(_id, _records)
            self._save_as_json(pt_records)

    @staticmethod
    def _save_as_json(pt_records: PatientRecords):
        if not os.path.exists(PCC_PT_RECORD_DIR):
            os.makedirs(PCC_PT_RECORD_DIR)
        filename = os.path.join(PCC_PT_RECORD_DIR, f"{pt_records.id}_pt_record.txt")

        def encode_records(obj):
            if isinstance(obj, PatientRecords):
                return {
                    '__pt_record__': True,
                    **vars(obj),
                }
            else:
                return getattr(obj, '__dict__', str(obj))

        jsn_str = json.dumps(pt_records, default=encode_records, indent=4, sort_keys=True)
        with open(filename, 'w') as writer:
            writer.write(jsn_str)

    @staticmethod
    def _read_from_json(file: str) -> PatientRecords:
        def decode_record(dct):
            if "__pt_record__" in dct:
                def decode_single(d: dict) -> PatientRecord:
                    d['date'] = datetime.strptime(d['date'], "%Y-%m-%d")
                    return PatientRecord(**d)

                return PatientRecords(
                    id=dct['id'],
                    records=[decode_single(r) for r in dct['records']]
                )
            return dct

        filename = os.path.join(PCC_PT_RECORD_DIR, file)
        with open(filename, 'r') as reader:
            data = reader.read()
            pt_transcript = json.loads(data, object_hook=decode_record)
            return pt_transcript

    def _get_doc_ids(self):
        if not os.path.exists(PCC_PT_RECORD_DIR):
            print(f'cannot find {PCC_PT_RECORD_DIR}')
            raise FileNotFoundError

        files = os.listdir(PCC_PT_RECORD_DIR)
        txt_docs = list(filter(lambda _file: _file.find('.txt') > 0, files))

        def _extract_id_from_name(filename):
            return re.search(r"^(?P<_id>.+)_pt_record\.txt", filename).group('_id')

        self.doc_ids = [_extract_id_from_name(filename) for filename in txt_docs]
        self.doc_ids.sort()

    @lru_cache(maxsize=128, typed=False)
    def get(self, record_id):
        if record_id in self.doc_ids:
            target = f"{record_id}_pt_record.txt"
            return self._read_from_json(target)
        else:
            logging.warning(f"{record_id} does not have PC record document.")


if __name__ == '__main__':
    pt_parser = PatientRecordParser()
